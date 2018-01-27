#! python3

import docx


def create_doc(doc_title):
    doc = docx.Document()
    doc.add_heading(doc_title, 0)
    doc.add_heading('The tale of Cassidy\'s first Python-generated .docx file', 1)
    doc.add_heading('January 27, 2018', 4)
    doc.add_paragraph('This is my first paragraph.')
    doc.add_paragraph('This is yet another paragraph.')
    doc.save('cassidysFirstDocx.docx')

create_doc('Cassidy\'s first Document')
